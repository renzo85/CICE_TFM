from flask import Flask
from flask import g, request, send_file
import re
from flask_cors import CORS
from selenium import webdriver
from selenium.common.exceptions import *
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as EC
import json, requests, ssl, os, logging, getpass, sys
from chromedriver_py import binary_path
from datetime import datetime
import time

from json import decoder
import docx
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pil import Image
# import csv
# import pathlib
#from PIL import Image

from docx.enum.style import WD_STYLE_TYPE

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://127.0.0.1:5500"}})

@app.route("/cv")

def cv():

    def GetCredentials():
        credentials = {}
        credentials['linkedin_user'] = input("Please input your linkedin email: ")
        credentials['linkedin_pwd'] = input("Please input your linkedin pwd: ")
        #credentials['linkedin_pwd'] = getpass.getpass("Please input your linkedin pwd (cursor wont move while typing): ")
        #credentials['linkedin_user'] = "rdcv54@gmail.com"
        #credentials['linkedin_pwd'] = "ncc1701d"
        return credentials

    def open_lkdinweb(credentials, driver):

        driver.get(r"https://www.linkedin.com/home")
        
        try:
            WebDriverWait(driver, 10).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[1]/div/section/div/div[2]/button[2]')))
            element = driver.find_element_by_xpath('/html/body/div[1]/div/section/div/div[2]/button[2]')
            element.click()

        except NoSuchElementException:
            pass
        
        try:
            #WebDriverWait(driver, 10).until(EC.visibility_of_element_located((By.ID,'session_key')))
            #WebDriverWait(driver, 10).until(EC.visibility_of_element_located((By.XPATH,'/html/body/main/section[1]/div/form/section/div[2]/div[1]/input')))
            WebDriverWait(driver, 10).until(EC.visibility_of_element_located((By.XPATH,'/html/body/main/section[1]/div/div/form/div[2]/div[1]/input')))                                                                             
            elementa = driver.find_element_by_xpath('/html/body/main/section[1]/div/div/form/div[2]/div[1]/input')
            elementa.send_keys(credentials['linkedin_user'])

        except:
            WebDriverWait(driver, 10).until(EC.visibility_of_element_located((By.XPATH,'/html/body/main/section[1]/div/form/section/div[2]/div[1]/input')))                                                                             
            elementa = driver.find_element_by_xpath('/html/body/main/section[1]/div/form/section/div[2]/div[1]/input')
            elementa.send_keys(credentials['linkedin_user'])

        try:
            WebDriverWait(driver, 10).until(EC.visibility_of_element_located((By.XPATH,'/html/body/main/section[1]/div/div/form/div[2]/div[2]/input')))
            #elementb = driver.find_element_by_id('session_password')
            elementb = driver.find_element_by_xpath('/html/body/main/section[1]/div/div/form/div[2]/div[2]/input')
            elementb.send_keys(credentials['linkedin_pwd'])
        
        except:
            WebDriverWait(driver, 10).until(EC.visibility_of_element_located((By.XPATH,'/html/body/main/section[1]/div/form/section/div[2]/div[2]/input')))
            elementb = driver.find_element_by_xpath('/html/body/main/section[1]/div/form/section/div[2]/div[2]/input')
            elementb.send_keys(credentials['linkedin_pwd'])
                                                                                                                        
        try:
            #elementc = driver.find_element_by_xpath('//*[@id="main-content"]/section[1]/div[2]/form/button')
            elementc = driver.find_element_by_xpath('/html/body/main/section[1]/div/div/form/button')
            elementc.click()
            time.sleep(2)

        except:
            elementc = driver.find_element_by_xpath('/html/body/main/section[1]/div/form/section/button')
            elementc.click()
            time.sleep(2)

        try:
            WebDriverWait(driver, 30).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div[2]/div/div/div/div[1]/div[1]/a/div[2]')))
            
        except:
            WebDriverWait(driver, 30).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[7]/div[3]/div/div/div[2]/div/div/div/div[1]/div[1]/a/div[2]')))
        
        #return "login_success"

    def open_lkdinprf(driver):

        #click on profile information
        try:
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div[2]/div/div/div/div[1]/div[1]/a/div[2]')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div[2]/div/div/div/div[1]/div[1]/a/div[2]')
            element.click()
        
        except:
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[7]/div[3]/div/div/div[2]/div/div/div/div[1]/div[1]/a/div[2]')))
            element = driver.find_element_by_xpath('/html/body/div[7]/div[3]/div/div/div[2]/div/div/div/div[1]/div[1]/a/div[2]')
            element.click()

    def extract_picture(driver):

        #download picture
        try:
            with open('Logo.png', 'wb') as file:
                WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/section[1]/div[2]/div[1]/div[1]/div/div/button/img')))
                element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/section[1]/div[2]/div[1]/div[1]/div/div/button/img')
                file.write(element.screenshot_as_png)
                                                        
        except:
            with open('Logo.png', 'wb') as file:
                WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[7]/div[3]/div/div/div/div/div[3]/div/div/main/div/section[1]/div[2]/div[1]/div[1]/div/div/button/img')))
                element = driver.find_element_by_xpath('/html/body/div[7]/div[3]/div/div/div/div/div[3]/div/div/main/div/section[1]/div[2]/div[1]/div[1]/div/div/button/img')
                file.write(element.screenshot_as_png)


    def extract_lkdinprf(driver):

        #click on "see more profile introduction button"
        try:
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[4]/section/div/span/button')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[4]/section/div/span/button')
            element.click()

        except:
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[7]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[4]/section/div/span/button')))
            element = driver.find_element_by_xpath('/html/body/div[7]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[4]/section/div/span/button')
            element.click()

        #extract profile introduction information text
        try:
            elementb = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[4]/section/div')
            profile = elementb.text
            
        except:
            elementb = driver.find_element_by_xpath('/html/body/div[7]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[4]/section/div')
            profile = elementb.text

        return profile


    def extract_name(driver):

        #extraction of name
        try:
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/section[1]/div[2]/div[2]/div/div[1]/h1')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/section[1]/div[2]/div[2]/div/div[1]/h1')
            name = element.text
            
        except:
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[7]/div[3]/div/div/div/div/div[3]/div/div/main/div/section[1]/div[2]/div[2]/div/div[1]/h1')))
            element = driver.find_element_by_xpath('/html/body/div[7]/div[3]/div/div/div/div/div[3]/div/div/main/div/section[1]/div[2]/div[2]/div/div[1]/h1')
            name = element.text

        return name


    def extract_subtittle(driver):

        #extraction of name's subtittle
        try:
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/section[1]/div[2]/div[2]/div[1]/div[2]')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/section[1]/div[2]/div[2]/div[1]/div[2]')
            subtittle = element.text
            
        except:
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[7]/div[3]/div/div/div/div/div[3]/div/div/main/div/section[1]/div[2]/div[2]/div[1]/div[2]')))
            element = driver.find_element_by_xpath('/html/body/div[7]/div[3]/div/div/div/div/div[3]/div/div/main/div/section[1]/div[2]/div[2]/div[1]/div[2]')
            subtittle = element.text

        return subtittle


    def extract_experience(driver):

        #click on "see more experience button"
        try:
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[1]/section/div/button')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[1]/section/div/button')
            element.click()
                    
        except:
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[7]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[1]/section/div/button')))
            element = driver.find_element_by_xpath('/html/body/div[7]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[1]/section/div/button')
            element.click()

        #extract whole list of job experience
        try:
            WebDriverWait(driver, 60).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[1]/section/ul')))
            #extract len from jobs
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[1]/section/ul')

            items = element.find_elements_by_tag_name("li")
            a = 1
            job = dict()

            while a <= len(items):

                #see more button, in order to extract job functions in a future version
                try:
                    WebDriverWait(driver, 2).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[1]/section/ul/li[{a}]/section/div/div[1]/div/div/span/button')))
                    element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[1]/section/ul/li[{a}]/section/div/div[1]/div/div/span/button')
                    element.click()

                except:
                    pass

                job [f"position{a}"] = driver.find_element_by_xpath(f"/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[1]/section/ul/li[{a}]/section/div/div[1]/a/div[2]/h3").text
                job [f"company{a}"] = driver.find_element_by_xpath(f"/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[1]/section/ul/li[{a}]/section/div/div[1]/a/div[2]/p[2]").text
                job [f"period{a}"] = driver.find_element_by_xpath(f"/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[1]/section/ul/li[{a}]/section/div/div[1]/a/div[2]/div/h4[1]/span[2]").text
                
                a += 1        
            return job

        except NoSuchElementException:
            pass


    def extract_education(driver):

        #click on "see more education button"
        try:
            WebDriverWait(driver, 60).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[2]/section/div/button')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[2]/section/div/button')
            element.click()
                    
        except NoSuchElementException:
            pass

        #extract whole list of education
        try:
            WebDriverWait(driver, 60).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[2]/section/ul')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[2]/section/ul')
            
            items = element.find_elements_by_tag_name("li")
            a = 1
            education = dict()
            
            while a < len(items):

                education [f"institution{a}"] = driver.find_element_by_xpath(f"/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[2]/section/ul/li[{a}]/div/div/div[1]/a/div[2]/div/h3").text
                education [f"program{a}"] = driver.find_element_by_xpath(f"/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[2]/section/ul/li[{a}]/div/div/div[1]/a/div[2]/div/p/span[2]").text
                education [f"period{a}"] = driver.find_element_by_xpath(f"/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[2]/span/div/section/div[2]/section/ul/li[{a}]/div/div/div[1]/a/div[2]/p/span[2]").text

                a += 1

            return education
                    
        except NoSuchElementException:
            pass


    def extract_skills(driver):

        #click on "see more skills button"
        try:
            WebDriverWait(driver, 60).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/button')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/button')
            element.click()
                    
        except NoSuchElementException:
            pass


        #extract whole list of skills
        try:
            WebDriverWait(driver, 60).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/ol')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/ol')
            
            items = element.find_elements_by_tag_name("li")
            a = 1
            skills = dict()
            b = len(items)/2

            while a <= b:
                                                                        
                skills [f"skill_name{a}"] = driver.find_element_by_xpath(f"/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/ol/li[{a}]/div/div/p").text
                skills [f"skill_value{a}"] = driver.find_element_by_xpath(f"/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/ol/li[{a}]/div/div/a/span[2]").text                                                  
                a += 1
                                                                        
            return skills
                    
        except NoSuchElementException:
            pass

    def extract_sector_knowledge(driver):

        #extract knowledge sector
        try:
            WebDriverWait(driver, 60).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/div[1]')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/div[1]')
            
            items = element.find_elements_by_tag_name("li")
            a = 1
            knowledge_sector = dict()

            while a <= len(items):

                knowledge_sector [f"knowsect_name{a}"] = driver.find_element_by_xpath(f"/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/div[1]/ol/li[{a}]/div/div/p").text
                a += 1

            return knowledge_sector
        
        except NoSuchElementException:
            pass       
                

    def extract_tools(driver):

        #extract tools
        try:
            WebDriverWait(driver, 60).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/div[2]')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/div[2]')
            
            items = element.find_elements_by_tag_name("li")
            a = 1
            tools = dict()

            while a <= len(items):
            
                tools [f"tool_name{a}"] = driver.find_element_by_xpath(f"/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/div[2]/ol/li[{a}]/div/div/p").text
                a += 1

            return tools

        except NoSuchElementException:
            pass

    def extract_pers_skills(driver):

        #extract personal skills
        try:
            WebDriverWait(driver, 60).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/div[3]')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/div[3]')
            
            items = element.find_elements_by_tag_name("li")
            a = 1
            perskills = dict()

            while a <= len(items):
                
                perskills [f"skill_name{a}"] = driver.find_element_by_xpath(f"/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/div[3]/ol/li[{a}]/div/div/p").text
                a += 1
            
            return perskills

        except NoSuchElementException:
            pass


    def extract_languages(driver):

        #extract languages
        try:
            WebDriverWait(driver, 60).until(EC.visibility_of_element_located((By.XPATH,'/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/div[4]')))
            element = driver.find_element_by_xpath('/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/div[4]')
            
            items = element.find_elements_by_tag_name("li")
            a = 1
            languages = dict()

            while a <= len(items):       
            
                languages [f"language{a}"] = driver.find_element_by_xpath(f"/html/body/div[6]/div[3]/div/div/div/div/div[3]/div/div/main/div/div[5]/div[3]/div/section/div[2]/div[4]/ol/li[{a}]/div/div/p").text
                a += 1
                
            return languages

        except NoSuchElementException:
            pass


#### Main #### 

    
    cvdict = {
        "name": "",
        "subtittle": "",
        "profile_intro": "",
        "job_experience": "",
        "education": "",
        "skills": "",
        "sector_know": "",
        "tools": "",
        "pers_skills": "",
        "languages": "",
        }

    credentials = GetCredentials()
    driver = webdriver.Chrome(executable_path=binary_path)
    driver.maximize_window()

    lknd_ln = open_lkdinweb(credentials, driver)

    lknd_pf_open = open_lkdinprf(driver)
    lknd_pic = extract_picture(driver)
    lknd_pf_extract = extract_lkdinprf(driver)
    cvdict['profile_intro'] = lknd_pf_extract
    lknd_nm = extract_name(driver)
    cvdict['name'] = lknd_nm
    lknd_st = extract_subtittle(driver)
    cvdict['subtittle'] = lknd_st
    lknd_jobs = extract_experience(driver)
    cvdict['job_experience'] = lknd_jobs
    lknd_edu = extract_education (driver)
    cvdict['education'] = lknd_edu
    lknd_skl = extract_skills (driver)
    cvdict['skills'] = lknd_skl
    lknd_sknow = extract_sector_knowledge (driver)
    cvdict['sector_know'] = lknd_sknow
    lknd_tools = extract_tools (driver)
    cvdict['tools'] = lknd_tools
    lknd_pskl = extract_pers_skills (driver)
    cvdict['pers_skills'] = lknd_pskl
    lknd_lg = extract_languages (driver)
    cvdict ['languages'] = lknd_lg



    cvfinal = json.dumps(cvdict, indent=4, ensure_ascii=False)
    print (cvfinal)


    with open ('cv.txt', 'w') as file:
        file.write(cvfinal)

    with open ('cv.json', 'w', encoding="utf8") as file:
        json.dump(cvdict, file,ensure_ascii=False, indent = 4 )
    
    driver.quit()

    
    # return send_file('cv.txt')


    with open ('cv.json', encoding='utf-8') as file:
        data = json.load(file)
        

    doc = docx.Document('template_0.docx')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:


                if "Texto0" in cell.text:
                    orig_text = cell.text
                    new_text = str.replace(orig_text,"Texto0", str(data['name']))
                    cell.text = new_text
                    paragraphs = cell.paragraphs
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Yu Gothic UI Light'
                            font.size = Pt(26)
                            font.bold = True
                            font.color.rgb= RGBColor(143, 117, 79)


                elif "Texto1" in cell.text:
                    orig_text = cell.text
                    new_text = str.replace(orig_text,"Texto1", str(data['subtittle']))
                    cell.text = new_text
                    paragraphs = cell.paragraphs
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Montserrat Light'
                            font.size = Pt(12)
                            font.bold = False
                            font.color.rgb= RGBColor(56, 95, 113)


                elif "Foto" in cell.text:
                    image = Image.open('Logo.png').convert('L')
                    image.save('Logobw.png')
                    orig_text = cell.text
                    new_text = str.replace(orig_text,"Foto","")
                    cell.text = new_text
                    cell.add_paragraph().add_run().add_picture('Logobw.png', width=Inches(1.72))
                

                elif "Texto2" in cell.text:
                    orig_text = cell.text
                    new_text = str.replace(orig_text,"Texto2", str(data['profile_intro']))
                    cell.text = new_text
                    paragraphs = cell.paragraphs
                    for paragraph in cell.paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Montserrat Light'
                            font.size = Pt(9)
                            font.bold = False
                            font.color.rgb= RGBColor(56, 95, 113)

                elif "Texto3" in cell.text:

                    orig_text = cell.text
                    new_text = str.replace(orig_text,"Texto3","")
                    cell.text = new_text

                    for value in data['pers_skills'].values():
                        cell.add_paragraph (str(value))
                    
                    paragraphs = cell.paragraphs
                    for paragraph in cell.paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        if paragraph.text == '' or paragraph.text == ' ' :
                            paragraph.style = doc.styles['Normal']
                        else:
                            paragraph.style = doc.styles['List Bullet']

                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Montserrat Light'
                            font.size = Pt(7)
                            font.bold = False
                            font.color.rgb= RGBColor(255, 255, 255)

                elif "Texto4" in cell.text:

                    orig_text = cell.text
                    new_text = str.replace(orig_text,"Texto4","")
                    cell.text = new_text

                    for key, value in data['skills'].items():
                        if "skill_name" in key:
                            cell.add_paragraph (str(value))

                    paragraphs = cell.paragraphs
                    for paragraph in cell.paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        if paragraph.text == '' or paragraph.text == ' ' :
                            paragraph.style = doc.styles['Normal']
                        else:
                            paragraph.style = doc.styles['List Bullet']

                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Montserrat Light'
                            font.size = Pt(7)
                            font.bold = False
                            font.color.rgb= RGBColor(255, 255, 255)

                elif "Texto5" in cell.text:

                    orig_text = cell.text
                    new_text = str.replace(orig_text,"Texto5","")
                    cell.text = new_text

                    for key, value in data['languages'].items():
                        if "language" in key:
                            cell.add_paragraph (str(value))

                    paragraphs = cell.paragraphs
                    for paragraph in cell.paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        if paragraph.text == '' or paragraph.text == ' ' :
                            paragraph.style = doc.styles['Normal']
                        else:
                            paragraph.style = doc.styles['List Bullet']

                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Montserrat Light'
                            font.size = Pt(7)
                            font.bold = False
                            font.color.rgb= RGBColor(255, 255, 255)


                elif "Texto6" in cell.text:

                    orig_text = cell.text
                    new_text = str.replace(orig_text,"Texto6","")
                    cell.text = new_text
                    counter_limit = 1

                    for key, value in data['sector_know'].items():

                        if counter_limit <= 5:

                            if "knowsect_name" in key:
                                cell.add_paragraph (str(value))
                            
                            counter_limit += 1
                        
                        else:
                            pass

                    paragraphs = cell.paragraphs
                    for paragraph in cell.paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        if paragraph.text == '' or paragraph.text == ' ' :
                            paragraph.style = doc.styles['Normal']
                        else:
                            paragraph.style = doc.styles['List Bullet']

                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Montserrat Light'
                            font.size = Pt(7)
                            font.bold = False
                            font.color.rgb= RGBColor(255, 255, 255)


                elif "Texto7" in cell.text:

                    orig_text = cell.text
                    new_text = str.replace(orig_text,"Texto7","")
                    cell.text = new_text
                    counter_limit = 1
                
                    for key, value in data['job_experience'].items():

                        if counter_limit <= 12:

                            if "position" in key:
                                cell.add_paragraph(str(value))
                                
                            elif "company" in key:
                                cell.add_paragraph(str(value))
                                                        
                            elif "period" in key:
                                cell.add_paragraph(str(value))
                                cell.add_paragraph(" ")
                            
                            counter_limit += 1
                            
                        else:
                            pass

                    paragraphs = cell.paragraphs
                    for paragraph in cell.paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        if paragraph.text == '' or paragraph.text == ' ' :
                            paragraph.style = doc.styles['Normal']
                        else:
                            paragraph.style = doc.styles['List Bullet']

                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Montserrat Light'
                            font.size = Pt(9)
                            font.bold = False
                            font.color.rgb= RGBColor(56, 95, 113)


                elif "Texto8" in cell.text:

                    orig_text = cell.text
                    new_text = str.replace(orig_text,"Texto8","")
                    cell.text = new_text
                    counter_limit = 1
                
                    for key, value in data['education'].items():

                        if counter_limit <= 12:

                            if "institution" in key:
                                cell.add_paragraph(str(value))
                                
                            elif "program" in key:
                                cell.add_paragraph(str(value))
                                                        
                            elif "period" in key:
                                cell.add_paragraph(str(value))
                                cell.add_paragraph(" ")
                            
                            counter_limit += 1
                            
                        else:
                            pass

                    paragraphs = cell.paragraphs
                    for paragraph in cell.paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        if paragraph.text == '' or paragraph.text == ' ' :
                            paragraph.style = doc.styles['Normal']
                        else:
                            paragraph.style = doc.styles['List Bullet']

                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Montserrat Light'
                            font.size = Pt(9)
                            font.bold = False
                            font.color.rgb= RGBColor(56, 95, 113)


    doc.save('template_0_saved.docx')
    
    return send_file('template_0_saved.docx')