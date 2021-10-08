from json import decoder
import docx
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
import pathlib
import json

with open ('cv.json', encoding='utf-8') as file:
    data = json.load(file)


doc = docx.Document('template_0.docx')

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:


            if "Texto1" in cell.text:
                orig_text = cell.text
                new_text = str.replace(orig_text,"Texto1", str(data['name']))
                cell.text = new_text
                paragraphs = cell.paragraphs
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Calibri'
                        font.size = Pt(26)
                        font.bold = True
                        font.color.rgb= RGBColor(204, 204, 0)

            elif "Foto" in cell.text:
                orig_text = cell.text
                new_text = str.replace(orig_text,"Foto","")
                cell.text = new_text
                cell.add_paragraph().add_run().add_picture('Logo.png', width=Inches(1.72))

            elif "Texto2" in cell.text:
                orig_text = cell.text
                new_text = str.replace(orig_text,"Texto2", str(data['profile_intro']))
                cell.text = new_text
                paragraphs = cell.paragraphs
                for paragraph in cell.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Calibri'
                        font.size = Pt(12)
                        font.bold = False
                        font.color.rgb= RGBColor(68, 114, 196)

            elif "Texto3" in cell.text:

                orig_text = cell.text
                new_text = str.replace(orig_text,"Texto3","")
                cell.text = new_text

                for value in data['pers_skills'].values():
                    cell.add_paragraph (str(value))
                
                paragraphs = cell.paragraphs
                for paragraph in cell.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Calibri'
                        font.size = Pt(9)
                        font.bold = False
                        font.color.rgb= RGBColor(255, 255, 255)

            elif "Texto4" in cell.text:

                orig_text = cell.text
                new_text = str.replace(orig_text,"Texto4","")
                cell.text = new_text

                for key, value in data['skills'].items():
                    if "skill_name" in key:
                        cell.add_paragraph (str(value))

                paragraphs = cell.paragraphs
                for paragraph in cell.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Calibri'
                        font.size = Pt(9)
                        font.bold = False
                        font.color.rgb= RGBColor(255, 255, 255)

            elif "Texto5" in cell.text:

                orig_text = cell.text
                new_text = str.replace(orig_text,"Texto5","")
                cell.text = new_text

                for key, value in data['languages'].items():
                    if "language" in key:
                        cell.add_paragraph (str(value))

                paragraphs = cell.paragraphs
                for paragraph in cell.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Calibri'
                        font.size = Pt(9)
                        font.bold = False
                        font.color.rgb= RGBColor(255, 255, 255)


            elif "Texto6" in cell.text:

                orig_text = cell.text
                new_text = str.replace(orig_text,"Texto6","")
                cell.text = new_text
                counter_limit = 1

                for key, value in data['sector_know'].items():

                    if counter_limit <= 5:

                        if "knowsect_name" in key:
                            cell.add_paragraph (str(value))
                        
                        counter_limit += 1
                    
                    else:
                        pass

                paragraphs = cell.paragraphs
                for paragraph in cell.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Calibri'
                        font.size = Pt(9)
                        font.bold = False
                        font.color.rgb= RGBColor(255, 255, 255)


            elif "Texto7" in cell.text:

                orig_text = cell.text
                new_text = str.replace(orig_text,"Texto7","")
                cell.text = new_text
                counter_limit = 1
            
                for key, value in data['job_experience'].items():

                    if counter_limit <= 12:

                        if "position" in key:
                            cell.add_paragraph(str(value))
                            
                        elif "company" in key:
                            cell.add_paragraph(str(value))
                                                       
                        elif "period" in key:
                            cell.add_paragraph(str(value))
                            cell.add_paragraph(" ")
                        
                        counter_limit += 1
                        
                    else:
                        pass

                paragraphs = cell.paragraphs
                for paragraph in cell.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Calibri'
                        font.size = Pt(12)
                        font.bold = False
                        font.color.rgb= RGBColor(68, 114, 196)


            elif "Texto8" in cell.text:

                orig_text = cell.text
                new_text = str.replace(orig_text,"Texto8","")
                cell.text = new_text
                counter_limit = 1
            
                for key, value in data['education'].items():

                    if counter_limit <= 12:

                        if "institution" in key:
                            cell.add_paragraph(str(value))
                            
                        elif "program" in key:
                            cell.add_paragraph(str(value))
                                                       
                        elif "period" in key:
                            cell.add_paragraph(str(value))
                            cell.add_paragraph(" ")
                        
                        counter_limit += 1
                        
                    else:
                        pass

                paragraphs = cell.paragraphs
                for paragraph in cell.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Calibri'
                        font.size = Pt(12)
                        font.bold = False
                        font.color.rgb= RGBColor(68, 114, 196)





doc.save('template_0_saved.docx')

   